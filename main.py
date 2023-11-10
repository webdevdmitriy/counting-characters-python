import docx, os, pypandoc, re, glob


# Удаление всех файлов из папки txt
for file in glob.glob("txt/*"):
    os.remove(file)

# Получаем все файлы ворд из папки docx. Преобразуем в txt и кладем в папку txt

file_uncounted = open("result/uncounted.txt", "w")
file_uncounted.write("Файлы, в которых символы не подсчитались:\n")


for filename in os.listdir("docx"):
    path_input = f"docx/{filename}"
    doc = docx.Document(path_input)

    # Имя файла или заголовок в документе
    # name = doc.paragraphs[0].text
    name = filename

    path_output = f"txt/{name}.txt"

    try:
        pypandoc.convert_file(path_input, "plain", outputfile=path_output)
    except Exception as error:
        file = open("result/uncounted.txt", "w")
        file.write(f"{filename}\n")
        print("Ошибка:", error)

file_uncounted.close()

# Подсчитываем кол-во символов в файлах
# Результаты заносим в файлы txt и word

doc = docx.Document()
count_files = len(os.listdir("txt"))

table = doc.add_table(rows=count_files, cols=3)
table.style = "Table Grid"


file_result = open("result/result.txt", "w")
row = 0


for filename in os.listdir("txt"):
    with open(os.path.join("txt", filename), "r", encoding="utf-8") as f:
        data = f.read()
        data_format = re.sub(r"\s\s+|\n|-", "", data)
        number_of_characters = len(data_format)
        result = f"{filename} {number_of_characters}\n"
        file_result.write(result)

        table.cell(row, 0).text = str(row + 1)
        table.cell(row, 1).text = re.sub(r"\.docx|\.txt", "", filename)
        table.cell(row, 2).text = str(number_of_characters)
        row += 1

file_result.close()
doc.save("result/result.docx")
